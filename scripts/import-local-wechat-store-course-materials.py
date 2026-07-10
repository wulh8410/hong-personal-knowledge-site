"""Import the local WeChat Store course pack into the static course library.

Each source item has a DOCX transcript and a matching infographic. The site
intentionally imports them as text lessons: replay URLs can be added to the
frontmatter later when the recordings are available.
"""

from __future__ import annotations

import argparse
import re
import shutil
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
COURSE_DIR = ROOT / "content" / "courses"
IMAGE_DIR = ROOT / "public" / "images" / "courses"

MATERIALS = [
    {
        "docx": "微信小店「一起买」功能介绍：从入门到精通.docx",
        "image": "「一起买」功能介绍.png",
        "slug": "wechat-store-group-buy",
        "title": "一起买：把拼单能力变成成交动作",
        "description": "理解一起买的核心机制、适用商品与活动设计，减少拼单链路中的流失。",
        "image_name": "wechat-store-group-buy.png",
    },
    {
        "docx": "微信小店「小店会员」功能使用指南：从入门到精通.docx",
        "image": "「小店会员」功能使用指南.png",
        "slug": "wechat-store-membership",
        "title": "小店会员：搭建复购与用户运营基础",
        "description": "从会员能力、权益配置到日常运营，建立可持续的复购机制。",
        "image_name": "wechat-store-membership.png",
    },
    {
        "docx": "微信小店「推荐场景」全面解读：规则、应用与运营指南.docx",
        "image": "「推荐场景」全面解读.png",
        "slug": "wechat-store-recommendation-scenes",
        "title": "推荐场景：让商品进入更多自然流量入口",
        "description": "梳理推荐场景的商品要求、运营判断与优化动作。",
        "image_name": "wechat-store-recommendation-scenes.png",
    },
    {
        "docx": "微信小店「搜索场景」全面解读：规则、运营与实战指南.docx",
        "image": "「搜索场景」全面解读.png",
        "slug": "wechat-store-search-scenes",
        "title": "搜索场景：用商品表达承接主动需求",
        "description": "从关键词、商品信息到搜索转化，建立可执行的优化路径。",
        "image_name": "wechat-store-search-scenes.png",
    },
    {
        "docx": "微信小店「电商罗盘」功能指南：用数据驱动生意增长 .docx",
        "image": "「电商罗盘」功能指南.png",
        "slug": "wechat-store-commerce-compass",
        "title": "电商罗盘：用数据驱动小店经营",
        "description": "看懂经营、商品、直播与市场数据，把异常识别变成具体动作。",
        "image_name": "wechat-store-commerce-compass.png",
    },
    {
        "docx": "微信小店「礼物营销」使用指南：从入门到精通.docx",
        "image": "「礼物营销」使用指南.png",
        "slug": "wechat-store-gift-marketing",
        "title": "礼物营销：设计社交送礼的增长场景",
        "description": "了解礼物营销的配置方式、活动场景和履约注意事项。",
        "image_name": "wechat-store-gift-marketing.png",
    },
    {
        "docx": "微信小店「营销优惠券」使用指南：从入门到精通.docx",
        "image": "「营销优惠券」使用指南.png",
        "slug": "wechat-store-marketing-coupons",
        "title": "营销优惠券：从发券到核销的经营闭环",
        "description": "拆解优惠券的配置、发放、核销和效果复盘方法。",
        "image_name": "wechat-store-marketing-coupons.png",
    },
    {
        "docx": "微信小店「评价抽奖」使用指南：从入门到精通.docx",
        "image": "「评价抽奖」使用指南.png",
        "slug": "wechat-store-review-lottery",
        "title": "评价抽奖：获取真实评价并带动复购",
        "description": "掌握评价抽奖的活动规则、设置步骤与风险边界。",
        "image_name": "wechat-store-review-lottery.png",
    },
]


def clean_text(text: str) -> str:
    text = " ".join(text.replace("\u00a0", " ").split())
    text = text.replace("**", "").replace("__", "")
    return text.strip("- ")


def to_markdown(docx_path: Path) -> str:
    document = Document(docx_path)
    lines: list[str] = []

    for index, paragraph in enumerate(document.paragraphs):
        text = clean_text(paragraph.text)
        if not text or re.fullmatch(r"[-─—_=]{3,}", text):
            continue
        if "点击图片可查看完整电子表格" in text:
            continue

        # The first line is the DOCX title and frontmatter already supplies it.
        if index == 0:
            continue

        if re.match(r"^[一二三四五六七八九十]+、", text):
            lines.extend([f"## {text}", ""])
        elif re.match(r"^[（(][一二三四五六七八九十0-9]+[）)]", text):
            lines.extend([f"### {text}", ""])
        elif len(text) <= 30 and not re.search(r"[。！？；]", text):
            lines.extend([f"### {text}", ""])
        else:
            lines.extend([text, ""])

    for table in document.tables:
        rows = [[clean_text(cell.text) for cell in row.cells] for row in table.rows]
        rows = [row for row in rows if any(row)]
        if len(rows) < 2:
            continue
        width = max(len(row) for row in rows)
        normalized = [row + [""] * (width - len(row)) for row in rows]
        lines.append("| " + " | ".join(normalized[0]) + " |")
        lines.append("| " + " | ".join(["---"] * width) + " |")
        for row in normalized[1:]:
            lines.append("| " + " | ".join(row) + " |")
        lines.append("")

    return "\n".join(lines).strip() + "\n"


def build_lesson(material: dict[str, str], order: int, source_dir: Path) -> None:
    transcript = to_markdown(source_dir / material["docx"])
    frontmatter = "\n".join(
        [
            "---",
            f'title: "{material["title"]}"',
            f'description: "{material["description"]}"',
            'track: "wechat-store"',
            f"order: {order}",
            'sourceIssue: "本地补充资料"',
            'sourceDate: "2026-07-10"',
            f'sourcePath: "资料/课程/微信小店/{material["docx"]}"',
            'replayUrl: ""',
            f'infographic: "/images/courses/{material["image_name"]}"',
            'status: "text"',
            "---",
            "",
        ]
    )
    (COURSE_DIR / f'{material["slug"]}.md').write_text(frontmatter + transcript, encoding="utf-8")
    shutil.copy2(source_dir / material["image"], IMAGE_DIR / material["image_name"])


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", required=True, type=Path)
    args = parser.parse_args()

    for order, material in enumerate(MATERIALS, start=9):
        build_lesson(material, order, args.source)
        print(f'Imported {material["slug"]}')


if __name__ == "__main__":
    main()
